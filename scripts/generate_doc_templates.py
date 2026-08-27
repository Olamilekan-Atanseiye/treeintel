# -*- coding: utf-8 -*-
"""
Generate placeholder .docx templates for species documentation.
=================================================
For every class in species_map.json that has a mapped filename but no
existing file in knowledge_docs/, this creates an empty .docx with the
correct filename and a consistent section structure (Description / Uses /
Ecology / Conservation Status) for you to fill in.

It never overwrites a file that already exists, and never invents content --
every section is left as a "TODO" placeholder for you to write or paste in.

Run from the project root:

    python scripts/generate_doc_templates.py
"""

import json
import sys
from pathlib import Path

from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "knowledge_docs"
MAP_PATH = BASE_DIR / "species_map.json"

SECTIONS = [
    ("Description", "TODO: general botanical description -- family, growth form, "
                     "leaves, bark, flowers/fruit, distinguishing features."),
    ("Uses", "TODO: ethnobotanical, traditional, medicinal, or commercial uses."),
    ("Ecology", "TODO: habitat, ecological role, distribution across tropical Africa."),
    ("Conservation Status", "TODO: IUCN status if known, threats, notes on scarcity or abundance."),
]


def make_template(species_name: str) -> Document:
    doc = Document()
    doc.add_heading(species_name, level=1)
    doc.add_paragraph(
        "This is a placeholder document generated automatically. "
        "Replace every TODO line below with real information before "
        "rebuilding the knowledge base -- the RAG system will only ever "
        "surface what's actually written here."
    )
    for heading, placeholder in SECTIONS:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(placeholder)
    return doc


def main():
    if not MAP_PATH.exists():
        sys.exit(f"species_map.json not found at {MAP_PATH}")

    species_map = json.loads(MAP_PATH.read_text(encoding="utf-8"))
    species_map.pop("_comment", None)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    created, skipped_existing, skipped_null = [], [], []

    for class_name, filename in species_map.items():
        if not filename:
            skipped_null.append(class_name)
            continue

        doc_path = DOCS_DIR / filename
        if doc_path.exists():
            skipped_existing.append(filename)
            continue

        doc = make_template(class_name)
        doc.save(str(doc_path))
        created.append(filename)

    print("=" * 60)
    print(f"Created {len(created)} new template(s) in {DOCS_DIR}")
    for f in created:
        print(f"  + {f}")

    if skipped_existing:
        print("-" * 60)
        print(f"Skipped {len(skipped_existing)} file(s) that already exist (not overwritten):")
        for f in skipped_existing:
            print(f"  = {f}")

    if skipped_null:
        print("-" * 60)
        print(f"Skipped {len(skipped_null)} class(es) mapped to null (no document expected):")
        for c in skipped_null:
            print(f"  - {c}")
    print("=" * 60)
    print("Next: open each new file in knowledge_docs/, replace the TODOs with "
          "real content, then run scripts/build_knowledge_base.py.")


if __name__ == "__main__":
    main()
